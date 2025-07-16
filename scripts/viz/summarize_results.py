import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import argparse
import os

# Border style definitions
STD_BORDER_SPEC = {"val": "single", "sz": "4", "color": "auto"}  # 0.5pt line
NO_BORDER_SPEC = {"val": "nil"}  # No border

# Model sizes in parameters
PLM_SIZES = {
    "prott5": 1_500_000_000,
    "prottucker": 1_500_000_000,
    "prostt5": 1_500_000_000,
    "clean": 650_000_000,
    "esm1b": 650_000_000,
    "esm2_8m": 8_000_000,
    "esm2_35m": 35_000_000,
    "esm2_150m": 150_000_000,
    "esm2_650m": 650_000_000,
    "esm2_3b": 3_000_000_000,
    "ankh_base": 450_000_000,
    "ankh_large": 1_150_000_000,
    "random_1024": 0,
}


def _set_cell_border(cell, **kwargs):
    """Set specific borders for a cell. kwargs can be top, bottom, left, right.
    Each kwarg should be a border specification dict like STD_BORDER_SPEC or NO_BORDER_SPEC.
    If a kwarg is None, that border is not touched.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for border_name_key, spec in kwargs.items():
        if spec is None:
            continue
        border_element_name = f"w:{border_name_key}"
        existing_border_el = tcBorders.find(qn(border_element_name))
        if existing_border_el is not None:
            tcBorders.remove(existing_border_el)
        border_el = OxmlElement(border_element_name)
        for att_name, att_val in spec.items():
            border_el.set(qn(f"w:{att_name}"), str(att_val))
        tcBorders.append(border_el)


def apply_cell_styling(paragraph, text_to_write, is_best, is_second_best):
    """Adds text to a paragraph and applies bold/underline/italic styling for Word."""
    run = paragraph.add_run(text_to_write)
    if is_best:
        run.bold = True
        run.underline = True
    elif is_second_best:
        run.bold = False
        run.underline = True
        run.italic = True


def create_word_metric_tables(
    csv_filepath, output_docx_filename="metrics_summary.docx"
):
    """
    Generates a Word document with styled tables summarizing model metrics.
    - Header row text is bold and centered. "Parameter" header is bold (not italic).
    - Header row has a bottom border.
    - First column data (Parameter names) is italic and centered.
    - Odd data rows have a background fill of F2F2F2 (excluding the first column).
    - All data cell text (Embeddings, metric values) is centered.
    - Horizontal borders only appear after each parameter section.
    - SE values are not displayed.
    For each metric (e.g., Pearson R2), a table is created where:
    - Rows are grouped by Parameter, then by Embedding.
    - Columns represent different Model Types.
    - Cell values are the metric scores, rounded to two decimal places.
    - For each Parameter and Model Type, the best and second-best performing
      Embeddings (based on the current metric) are highlighted.
    - Parameter cells in the first column are vertically merged and centered.
    """
    try:
        df = pd.read_csv(csv_filepath)
    except FileNotFoundError:
        print(f"Error: File not found at {csv_filepath}")
        return
    except Exception as e:
        print(f"An error occurred while reading the CSV: {e}")
        return

    doc = Document()
    doc.add_heading("Metrics Summary Tables", level=0)

    metrics_to_tabulate = ["Pearson R2", "MAE", "Spearman", "R2"]
    # Sort model types by their parameter size
    model_type_columns = sorted(
        df["Model Type"].unique(),
        key=lambda x: PLM_SIZES.get(
            x.lower(), float("inf")
        ),  # Use model size for sorting, default to inf if unknown
    )
    parameter_names = sorted(df["Parameter"].unique())

    table_style_to_apply = "Table Grid"

    for metric_name in metrics_to_tabulate:
        doc.add_heading(f"Metric: {metric_name}", level=1)

        header_labels = ["Parameter", "Embedding"] + model_type_columns
        table = doc.add_table(rows=1, cols=len(header_labels))
        table.style = table_style_to_apply

        hdr_cells = table.rows[0].cells
        for i, label in enumerate(header_labels):
            cell = hdr_cells[i]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(label)
            run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_border(
                cell, bottom=STD_BORDER_SPEC
            )  # Ensure header bottom border
            # Top, Left, Right from Table Grid style

        tblPr = table._element.xpath("w:tblPr")[0]
        tblLook = tblPr.find(qn("w:tblLook"))
        if tblLook is None:
            tblLook = OxmlElement("w:tblLook")
            tblPr.append(tblLook)
        tblLook.set(qn("w:firstRow"), "1")  # Header row formatting active
        tblLook.set(qn("w:noHBand"), "0")  # Horizontal banding active (for row shading)
        tblLook.set(qn("w:noVBand"), "1")  # No vertical banding (column shading)

        styling_refs = {}
        for mt_col in model_type_columns:
            for param_group in parameter_names:
                relevant_data = df[
                    (df["Model Type"] == mt_col) & (df["Parameter"] == param_group)
                ]
                if relevant_data.empty or metric_name not in relevant_data.columns:
                    styling_refs[(mt_col, param_group)] = {
                        "best": None,
                        "second_best": None,
                    }
                    continue

                scores = relevant_data[metric_name].dropna().astype(float).tolist()
                best_s, second_best_s = None, None
                if scores:
                    unique_scores = sorted(list(set(scores)))
                    if metric_name != "MAE":
                        unique_scores.reverse()
                    best_s = unique_scores[0] if len(unique_scores) > 0 else None
                    if len(unique_scores) > 1:
                        second_best_s = unique_scores[1]
                styling_refs[(mt_col, param_group)] = {
                    "best": best_s,
                    "second_best": second_best_s,
                }

        current_table_abs_row_idx = 0
        data_row_counter_for_banding = 0

        for param_name in parameter_names:
            # Get embeddings for this parameter and sort them by model size
            embeddings_for_param = sorted(
                df[df["Parameter"] == param_name]["Embedding"].unique(),
                key=lambda x: PLM_SIZES.get(
                    x.lower(), float("inf")
                ),  # Sort by model size
            )
            if not embeddings_for_param:
                continue

            param_block_start_row_abs_idx = current_table_abs_row_idx + 1

            for embed_idx, embed_name in enumerate(embeddings_for_param):
                row_cells = table.add_row().cells
                current_table_abs_row_idx += 1
                is_last_embedding_in_block = embed_idx == len(embeddings_for_param) - 1

                if data_row_counter_for_banding % 2 == 0:
                    band_color = "F2F2F2"
                    for cell_to_band in row_cells[1:]:
                        tcPr_band = cell_to_band._tc.get_or_add_tcPr()
                        shading_elm_band = OxmlElement("w:shd")
                        shading_elm_band.set(qn("w:val"), "clear")
                        shading_elm_band.set(qn("w:fill"), band_color)
                        tcPr_band.append(shading_elm_band)

                data_row_counter_for_banding += 1

                if embed_idx == 0:
                    pass

                row_cells[1].text = str(embed_name)

                for cell_in_row in row_cells:
                    _set_cell_border(
                        cell_in_row,
                        top=NO_BORDER_SPEC,
                        bottom=STD_BORDER_SPEC
                        if is_last_embedding_in_block
                        else NO_BORDER_SPEC,
                        # Left & Right borders will come from "Table Grid" style
                    )

                for model_col_idx, mt_col_name in enumerate(model_type_columns):
                    cell_data_df = df[
                        (df["Parameter"] == param_name)
                        & (df["Embedding"] == embed_name)
                        & (df["Model Type"] == mt_col_name)
                    ]

                    val_str = ""
                    numeric_val = None
                    if (
                        not cell_data_df.empty
                        and metric_name in cell_data_df.columns
                        and pd.notna(cell_data_df.iloc[0][metric_name])
                    ):
                        original_val = cell_data_df.iloc[0][metric_name]
                        try:
                            f_val = float(original_val)
                            val_str = f"{f_val:.2f}"
                            numeric_val = f_val
                        except (ValueError, TypeError):
                            val_str = str(original_val)

                    ref_key = (mt_col_name, param_name)
                    best_score = styling_refs.get(ref_key, {}).get("best")
                    second_best_score = styling_refs.get(ref_key, {}).get("second_best")

                    is_best = (
                        numeric_val is not None
                        and best_score is not None
                        and numeric_val == best_score
                    )
                    is_second_best = False
                    if (
                        not is_best
                        and numeric_val is not None
                        and second_best_score is not None
                        and numeric_val == second_best_score
                    ):
                        is_second_best = True

                    table_col_idx = model_col_idx + 2
                    p = row_cells[table_col_idx].paragraphs[0]
                    p.clear()
                    apply_cell_styling(p, val_str, is_best, is_second_best)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if len(embeddings_for_param) > 1:
                merged_cell = table.cell(param_block_start_row_abs_idx, 0).merge(
                    table.cell(current_table_abs_row_idx, 0)
                )
                p_merged = merged_cell.paragraphs[0]
                p_merged.clear()
                run_merged = p_merged.add_run(str(param_name))
                run_merged.italic = True

                p_merged.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_merged.paragraph_format.space_before = Pt(0)
                p_merged.paragraph_format.space_after = Pt(0)
                p_merged.paragraph_format.line_spacing = 1.0
                merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            elif len(embeddings_for_param) == 1:
                cell_to_format = table.cell(param_block_start_row_abs_idx, 0)
                p_single = cell_to_format.paragraphs[0]
                p_single.clear()
                run_single = p_single.add_run(str(param_name))
                run_single.italic = True

                p_single.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_single.paragraph_format.space_before = Pt(0)
                p_single.paragraph_format.space_after = Pt(0)
                p_single.paragraph_format.line_spacing = 1.0
                cell_to_format.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        doc.add_paragraph()

    output_dir = os.path.dirname(csv_filepath)
    # If csv_filepath is just a filename (no dir part), output_dir will be '',
    # and os.path.join will correctly place the output file in the CWD.
    actual_output_docx_filename = os.path.join(output_dir, "metrics_summary.docx")

    doc.save(actual_output_docx_filename)
    print(f"Word document '{actual_output_docx_filename}' created successfully.")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate a Word document with styled tables summarizing model metrics from a CSV file."
    )
    parser.add_argument(
        "--csv_filepath",
        type=str,
        default="out/summary_plot/parsed_metrics_all.csv",
        help="Path to the input CSV file. (Default: out/summary_plot/parsed_metrics_all.csv)",
    )
    args = parser.parse_args()

    create_word_metric_tables(args.csv_filepath)
